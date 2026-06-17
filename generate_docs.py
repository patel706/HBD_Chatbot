# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2)

# ── Colour palette ─────────────────────────────────────────────────────────────
BRAND_BLUE  = RGBColor(0x1E, 0x40, 0xAF)   # #1E40AF
BRAND_TEAL  = RGBColor(0x0F, 0x76, 0x6E)   # #0F766E
DARK_TEXT   = RGBColor(0x1F, 0x29, 0x37)   # #1F2937
ACCENT_GREY = RGBColor(0x6B, 0x72, 0x80)   # #6B7280

def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def heading(level, text, color=BRAND_BLUE):
    h = doc.add_heading(level=level)
    h.clear()
    run = h.add_run(text)
    if level == 1:
        set_font(run, size=20, bold=True, color=color)
    elif level == 2:
        set_font(run, size=14, bold=True, color=color)
    else:
        set_font(run, size=12, bold=True, color=BRAND_TEAL)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after  = Pt(6)
    return h

def para(text, bold=False, size=11, color=DARK_TEXT, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=10.5, color=DARK_TEXT)
    p.paragraph_format.left_indent  = Cm(0.8 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    return p

def separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '93C5FD')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E40AF')
        tcPr.append(shd)
    # Data rows
    for row_data in rows:
        row = tbl.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for run in row.cells[i].paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run("🤖 CityHangAround AI Chatbot")
run.font.name = "Calibri"
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = BRAND_BLUE
cover_title.paragraph_format.space_before = Pt(80)
cover_title.paragraph_format.space_after  = Pt(10)

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = cover_sub.add_run("Complete Project Documentation")
run2.font.name = "Calibri"
run2.font.size = Pt(18)
run2.font.color.rgb = BRAND_TEAL
cover_sub.paragraph_format.space_after = Pt(6)

cover_tag = doc.add_paragraph()
cover_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = cover_tag.add_run("Functionalities · Architecture · Workflow · Tech Stack")
run3.font.name = "Calibri"
run3.font.size = Pt(12)
run3.font.italic = True
run3.font.color.rgb = ACCENT_GREY
cover_tag.paragraph_format.space_after = Pt(60)

cover_date = doc.add_paragraph()
cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = cover_date.add_run(f"Prepared by: Honeybee Digital\nDate: {datetime.date.today().strftime('%B %d, %Y')}")
run4.font.name = "Calibri"
run4.font.size = Pt(11)
run4.font.color.rgb = DARK_TEXT

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "1. Project Overview")
separator()
para(
    "CityHangAround AI Chatbot is a full-stack, multilingual, AI-powered chatbot "
    "built for a local business discovery and management platform. It lets everyday "
    "users search for nearby businesses and lets business owners register, verify, "
    "and manage their listings — all through a conversational chat interface embedded "
    "in a modern web application.",
    size=11
)
para("")
heading(3, "Key Goals")
bullet("Allow users to search businesses by category, name, or location")
bullet("Let business owners register and manage their listings via chat")
bullet("Support 16 Indian regional languages with dynamic auto-translation")
bullet("Provide secure authentication via email OTP or direct phone login")
bullet("Enable product and deal management within the chat")
bullet("Work fully offline using a local SQLite database and CSV data store")

# ══════════════════════════════════════════════════════════════════════════════
#  2. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "2. Technology Stack")
separator()

heading(2, "2.1 Frontend")
add_table(
    ["Technology", "Purpose", "Version"],
    [
        ["React.js", "Main UI framework for the chatbot widget", "18.x"],
        ["Vite", "Build tool and local dev server", "Latest"],
        ["Tailwind CSS", "Utility-first CSS styling", "3.x"],
        ["Lucide React", "Icon library for UI elements", "Latest"],
        ["JavaScript (ES6+)", "Core programming language", "—"],
    ]
)

heading(2, "2.2 Backend")
add_table(
    ["Technology", "Purpose", "Version"],
    [
        ["Python 3.12", "Core backend language", "3.12"],
        ["FastAPI", "REST API framework", "Latest"],
        ["Uvicorn", "ASGI web server with --reload support", "Latest"],
        ["SQLite", "Local relational database", "Built-in"],
        ["SQLite3 (Python)", "Database driver for queries", "Built-in"],
        ["python-dotenv", "Environment variable management", "Latest"],
        ["smtplib / email", "Email OTP sending via SMTP", "Built-in"],
        ["difflib", "Fuzzy matching for typo correction", "Built-in"],
    ]
)

heading(2, "2.3 AI / LLM Integration")
add_table(
    ["Component", "Role"],
    [
        ["LLM Client (llm_client.py)", "Calls a configured Large Language Model for intent classification, SQL generation, and responses"],
        ["assistant_manager.py", "Manages system prompts, intent classification, greeting detection, and multilingual responses"],
        ["text_to_sql.py", "Converts natural language queries to safe, read-only SQL for the database"],
        ["fast_result.py", "Returns quick, direct answers for common FAQs without LLM call overhead"],
        ["search_online.py", "Searches online for business listings when no local results exist (SEARCH_BUSINESS only)"],
    ]
)

heading(2, "2.4 Data Storage")
add_table(
    ["Storage", "File", "Purpose"],
    [
        ["SQLite DB", "google_map_data.db", "Primary store for all business, product, and deal records"],
        ["CSV File", "g_map_master_table_sample.csv", "Secondary data source / backup master list (14-column schema)"],
        ["Static Files", "backend/static/uploads/", "Product image uploads served via FastAPI static mount"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  3. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "3. System Architecture")
separator()
para("The chatbot follows a client-server architecture with the following layers:")
para("")
heading(3, "Frontend Layer  (React + Vite)  →  Port 5173")
bullet("ChatWidget.jsx — The main chatbot component (1000+ line React component)")
bullet("MessageItem.jsx — Renders individual messages (text, suggestions, cards, products, deals)")
bullet("LoginPopup.jsx — Handles user authentication (Phone or Email OTP login)")
bullet("QuickActions.jsx — Renders dynamic action buttons based on user state")
bullet("Translations.js — Full translations for 16 languages; uses a Proxy for automatic English fallback")
bullet("Languages.jsx — Defines all supported Indian language codes and display names")

para("")
heading(3, "Backend Layer  (FastAPI)  →  Port 5000")
bullet("api.py — Main REST API with all endpoints, business logic, and localization engine")
bullet("assistant_manager.py — NLP intent classification and multilingual response generation")
bullet("text_to_sql.py — Safe natural language → SQL converter (READ-ONLY queries)")
bullet("business_by_phone.py — Phone and email lookup functions for the database")
bullet("business_update.py — Handles field-level business record updates")
bullet("fast_result.py — FAQ fast-path handler")
bullet("search_online.py — Online fallback search (triggered only for SEARCH_BUSINESS)")
bullet("llm_client.py — Thin wrapper around the LLM API")
bullet("models.py — LLM model name configuration")

para("")
heading(3, "Startup Scripts")
bullet("START_FULL_APP.bat — Starts both frontend (npm run dev) and backend simultaneously")
bullet("START_SERVER.bat — Starts only the FastAPI backend with --reload flag")

# ══════════════════════════════════════════════════════════════════════════════
#  4. DATABASE SCHEMA
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "4. Database Schema")
separator()

heading(2, "4.1  google_maps_listings  (Main Business Table)")
add_table(
    ["Column", "Type", "Description"],
    [
        ["id", "INTEGER PK", "Auto-incremented unique business ID"],
        ["name", "TEXT", "Business name"],
        ["address", "TEXT", "Full street address"],
        ["phone_number", "TEXT", "Contact phone number (10 digits)"],
        ["email", "TEXT", "Business owner email address"],
        ["category", "TEXT", "Business category (e.g. Resort, Hospital)"],
        ["subcategory", "TEXT", "Sub-category"],
        ["website", "TEXT", "Website URL"],
        ["area", "TEXT", "Local area / locality name"],
        ["city", "TEXT", "City name"],
        ["state", "TEXT", "State name"],
        ["reviews_count", "INTEGER", "Total Google review count"],
        ["reviews_average", "REAL", "Average star rating"],
        ["created_at", "TEXT", "ISO timestamp of record creation"],
    ]
)

heading(2, "4.2  products  (Business Products)")
add_table(
    ["Column", "Type", "Description"],
    [
        ["id", "INTEGER PK", "Auto product ID"],
        ["business_id", "INTEGER FK", "References google_maps_listings.id"],
        ["name", "TEXT", "Product name"],
        ["price", "REAL", "Selling price"],
        ["category", "TEXT", "Product category"],
        ["description", "TEXT", "Short description"],
        ["image_url", "TEXT", "URL to the uploaded product image"],
    ]
)

heading(2, "4.3  deals  (Business Deals / Offers)")
add_table(
    ["Column", "Type", "Description"],
    [
        ["id", "INTEGER PK", "Auto deal ID"],
        ["business_id", "INTEGER FK", "References google_maps_listings.id"],
        ["title", "TEXT", "Deal/offer title"],
        ["discount_pct", "INTEGER", "Discount percentage (e.g. 20)"],
        ["expiry_date", "TEXT", "Deal expiry date (YYYY-MM-DD)"],
        ["description", "TEXT", "Detailed offer description"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  5. API ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "5. API Endpoints")
separator()
add_table(
    ["Method", "Endpoint", "Description"],
    [
        ["POST", "/query", "Main chat endpoint — processes all user messages, returns typed responses"],
        ["POST", "/login", "Logs in a user by phone or email; returns matching business records"],
        ["POST", "/business", "Registers a new business listing (phone, email, name, category, etc.)"],
        ["PUT", "/business/{id}", "Updates a specific field of a business record"],
        ["GET", "/business/by-phone/{phone}", "Returns businesses linked to a phone number"],
        ["GET", "/business/{id}/suggestions", "Returns AI-generated field improvement suggestions"],
        ["POST", "/api/send-otp", "Sends a 6-digit email OTP via SMTP for login or registration"],
        ["POST", "/api/verify-otp", "Verifies the submitted OTP code"],
        ["POST", "/api/products", "Adds a new product to a business"],
        ["DELETE", "/api/products/{id}", "Removes a product"],
        ["POST", "/api/deals", "Adds a new deal/offer to a business"],
        ["DELETE", "/api/deals/{id}", "Removes a deal"],
        ["POST", "/api/upload", "Uploads a product image; returns static URL"],
        ["GET", "/api/business/search-name", "Searches businesses by name keyword"],
        ["GET", "/api/business/search-address", "Searches businesses by address/area keyword"],
        ["GET", "/api/smart-suggestions", "Returns AI-powered action suggestions for the chat"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  6. CORE FUNCTIONALITIES
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "6. Core Functionalities")
separator()

heading(2, "6.1 User Authentication")
bullet("Phone Login — User enters a 10-digit mobile number; direct login without OTP or reCAPTCHA")
bullet("Email OTP Login — User enters an email address; system sends a 6-digit OTP via SMTP; verified on submission")
bullet("Guest Mode — Users can search businesses and use the chatbot without logging in")
bullet("Session Management — On successful login, a session object is stored in React state containing phone, email, and businessId")
bullet("Auto-Detection — Logging in via email automatically links any business registered with that email")

heading(2, "6.2 Business Discovery (Search)")
bullet("Natural Language Search — Type anything like 'pizza in Surat' or 'best resort in Gujarat'")
bullet("Name Search — Explicit business name lookup using a dedicated search flow")
bullet("Address / Area Search — Find businesses by locality or street address")
bullet("AI SQL Generation — Natural language queries are converted to safe SQL via text_to_sql.py and run against the local database")
bullet("Online Fallback — If no local result is found AND the intent is SEARCH_BUSINESS, the system fetches online results via search_online.py")
bullet("Fuzzy Matching — Minor typos in typed commands (e.g. 'add prodcut') are auto-corrected using difflib's get_close_matches")

heading(2, "6.3 Business Registration (Add Wizard)")
bullet("Step-by-step guided wizard with 9 steps: Mobile Number → Email → OTP → Business Name → Category → Address → City → Area → State")
bullet("Real-time email OTP sent during registration to verify the business owner's email")
bullet("10-digit phone validation enforced on the phone step")
bullet("On completion: record saved to SQLite DB + appended to CSV master file")
bullet("Session auto-upgraded to BUSINESS type with the new businessId")

heading(2, "6.4 Business Profile Management")
bullet("'Show My Business' — Displays the logged-in user's full business profile in a structured, translated card")
bullet("'Update My Business' — Shows a list of all editable fields as one-column suggestion tiles")
bullet("Field Update Flow — User selects a field, types a new value, update is sent via PUT /business/{id}")
bullet("Supported update fields: Name, Category, Phone Number, Address, Area, City, State, Website")
bullet("Update success/ failure messages are shown in the currently selected language")

heading(2, "6.5 Product Management")
bullet("Add Product Wizard — Guided 5-step flow: Name → Price → Category → Description → Image Upload")
bullet("Product image upload support via file picker, stored in backend/static/uploads/")
bullet("'Manage Products' — Shows all products for the business as interactive cards")
bullet("Delete Product — Removes the product from the database instantly")
bullet("Products are tied to the business via business_id foreign key")

heading(2, "6.6 Deal / Offer Management")
bullet("Add Deal Wizard — Guided 4-step flow: Title → Discount % → Expiry Date → Description")
bullet("'Manage Deals' — Shows all active deals as interactive cards with delete option")
bullet("Percentage validation ensures only numeric values are accepted for discount_pct")
bullet("Deals are tied to the business via business_id foreign key")

heading(2, "6.7 Multilingual Support (16 Languages)")
bullet("Frontend: All UI strings, button labels, wizard prompts, and error messages are translated using Translations.js")
bullet("Translation Proxy: Any missing key automatically falls back to the English equivalent — no crashes")
bullet("Backend: All API response strings (profile headings, field labels, update prompts) use lang_fetch() with LLM fallback")
bullet("LLM Translation Fallback: If a string is not in the static dictionary, the LLM auto-translates it on demand and caches the result in memory")
bullet("Languages supported: English, Hindi, Gujarati, Telugu, Tamil, Marathi, Bengali, Kannada, Malayalam, Punjabi, Odia, Assamese, Urdu, Sanskrit, Kashmiri, Nepali")
bullet("Language selector is available in the chat header; changing language updates both UI and chat in real time")

heading(2, "6.8 Conversational AI (LLM Integration)")
bullet("Intent Classification — Every user message is classified into one of: FAQ, SEARCH_BUSINESS, BUSINESS_STATUS, BUSINESS_UPDATE, LOGIN, FAST_RESULT, UNKNOWN")
bullet("Greeting Detection — Pure regex-based, no LLM call needed for greetings (hi, hello, namaste, etc.)")
bullet("FAQ Fast Path — Common platform questions answered immediately without LLM")
bullet("AI Responses — For complex queries, the LLM generates a contextual response in the user's language")
bullet("Strict Security — The system prompt enforces READ-ONLY SQL, rejects prompt injection, and never exposes internal data")
bullet("Language-aware — The LLM is always instructed to respond in the user's selected language")

heading(2, "6.9 Typo / Error Tolerance")
bullet("Server-side fuzzy matching using difflib for command recognition (70% similarity threshold)")
bullet("Commands like 'add busines', 'manag deals', 'loign' are auto-corrected to the correct action")
bullet("Client-side 10-digit enforcement strips non-numeric characters from phone inputs in real time")
bullet("Email format validation before sending OTP to prevent wasted API calls")

# ══════════════════════════════════════════════════════════════════════════════
#  7. WORKFLOW / USER JOURNEY
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "7. Complete User Workflow")
separator()

heading(2, "7.1 Guest User — Search a Business")
for i, step in enumerate([
    "User opens the chatbot widget on the website",
    "Welcome message displayed in English (or selected language)",
    "User types a query e.g. 'best hotels in Surat' or clicks 'Find My Business'",
    "Backend receives the query at POST /query",
    "intent = classify_intent(query) → returns SEARCH_BUSINESS",
    "text_to_sql.py generates SQL → executes against google_map_data.db",
    "Results returned as a list of business cards",
    "If no local results → online search triggered via search_online.py",
    "User sees results in the chat with name, area, city, rating displayed",
], 1):
    bullet(f"Step {i}: {step}")

heading(2, "7.2 Business Owner — Register New Business")
for i, step in enumerate([
    "User clicks 'Add New Business' or types 'add business'",
    "ADD_WIZARD flow starts (9 guided steps)",
    "User enters mobile number → validated (10 digits)",
    "User enters email → OTP sent via SMTP",
    "User enters OTP → verified via POST /api/verify-otp",
    "User enters Business Name, Category, Address, City, Area, State",
    "On final step → POST /business called with all data",
    "Record saved to SQLite + appended to CSV with email field",
    "User session upgraded; business management panel appears",
], 1):
    bullet(f"Step {i}: {step}")

heading(2, "7.3 Business Owner — Update Business")
for i, step in enumerate([
    "User logs in with phone or email",
    "Session set with businessId",
    "User types 'update my business' or clicks button",
    "Backend returns list of all editable fields as suggestion tiles (1-column)",
    "User clicks a tile (e.g. 'Update Phone Number')",
    "Bot asks for new value in the selected language",
    "User types new value → validation run (10-digit check for phone)",
    "PUT /business/{id} called → database updated",
    "Success message shown in selected language",
], 1):
    bullet(f"Step {i}: {step}")

heading(2, "7.4 Login Flow — Email OTP")
for i, step in enumerate([
    "User clicks 'Login' in the chat header",
    "LoginPopup opens with Phone / Email toggle",
    "User selects 'Email' tab and enters email address",
    "POST /api/send-otp → 6-digit OTP sent via SMTP",
    "User enters OTP in the input field",
    "POST /api/verify-otp → verified",
    "POST /login called with email → businesses fetched",
    "Session set, chatbot transitions to business management view",
], 1):
    bullet(f"Step {i}: {step}")

# ══════════════════════════════════════════════════════════════════════════════
#  8. SECURITY
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "8. Security & Data Integrity")
separator()
bullet("All SQL queries generated by the AI are strictly READ-ONLY; DELETE, UPDATE, DROP, INSERT, ALTER are rejected")
bullet("Prompt injection attempts are detected and rejected by the LLM system prompt")
bullet("OTP codes expire and are single-use; stored temporarily in server memory")
bullet("Phone numbers are normalized (strip spaces, country codes) before database lookup")
bullet("Email login cross-references the email in the business record to prevent unauthorized access")
bullet("Business updates (PUT) require a valid businessId in the session — no unauthenticated edits possible")
bullet("No sensitive data (passwords, API keys) is exposed in chat responses")
bullet("CORS is configured with explicit origin whitelist in api.py")
bullet("reCAPTCHA and Firebase Phone Auth have been removed to streamline the flow; security is maintained via email OTP")

# ══════════════════════════════════════════════════════════════════════════════
#  9. FILE STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "9. Project File Structure")
separator()
for line in [
    "frontend 4/",
    "├── backend/",
    "│   ├── api.py                     ← Main FastAPI app + localization engine",
    "│   ├── assistant_manager.py       ← Intent classification + multilingual AI",
    "│   ├── business_by_phone.py       ← Phone/email database lookup",
    "│   ├── business_update.py         ← Field-level update logic",
    "│   ├── text_to_sql.py             ← NL → SQL converter",
    "│   ├── fast_result.py             ← FAQ fast-path handler",
    "│   ├── search_online.py           ← Online search fallback",
    "│   ├── llm_client.py              ← LLM API wrapper",
    "│   ├── models.py                  ← LLM model configuration",
    "│   ├── google_map_data.db         ← SQLite database",
    "│   ├── g_map_master_table_sample.csv  ← Master CSV (14 columns)",
    "│   ├── static/uploads/            ← Product image storage",
    "│   ├── START_SERVER.bat           ← Backend startup script",
    "│   └── .env                       ← Email credentials & config",
    "├── src/",
    "│   ├── components/",
    "│   │   ├── Chatwidget.jsx         ← Main chatbot component",
    "│   │   ├── MessageItem.jsx        ← Message rendering (all types)",
    "│   │   ├── LoginPopup.jsx         ← Authentication popup",
    "│   │   └── QuickActions.jsx       ← Action button panels",
    "│   ├── constants/",
    "│   │   ├── Translations.js        ← 16-language translation dictionary",
    "│   │   ├── Languages.jsx          ← Language codes and labels",
    "│   │   └── Suggestions.js         ← Static suggestion prompts",
    "│   └── firebase.js                ← Firebase config (retained for reference)",
    "├── START_FULL_APP.bat             ← Starts frontend + backend together",
    "└── package.json                   ← Frontend dependencies",
]:
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.name = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = DARK_TEXT
    p.paragraph_format.space_after = Pt(1)

# ══════════════════════════════════════════════════════════════════════════════
#  10. SUPPORTED LANGUAGES
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "10. Supported Languages")
separator()
add_table(
    ["Code", "Language", "Script"],
    [
        ["en", "English", "Latin"],
        ["hi", "Hindi", "Devanagari"],
        ["gu", "Gujarati", "Gujarati"],
        ["te", "Telugu", "Telugu"],
        ["ta", "Tamil", "Tamil"],
        ["mr", "Marathi", "Devanagari"],
        ["bn", "Bengali", "Bengali"],
        ["kn", "Kannada", "Kannada"],
        ["ml", "Malayalam", "Malayalam"],
        ["pa", "Punjabi", "Gurmukhi"],
        ["or", "Odia", "Odia"],
        ["as", "Assamese", "Bengali"],
        ["ur", "Urdu", "Nastaliq (Arabic)"],
        ["sa", "Sanskrit", "Devanagari"],
        ["ks", "Kashmiri", "Nastaliq (Arabic)"],
        ["ne", "Nepali", "Devanagari"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  11. HOW TO RUN
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "11. How to Run the Project")
separator()

heading(2, "Option A — Full App (Recommended)")
p = doc.add_paragraph()
r = p.add_run("Double-click:  START_FULL_APP.bat")
r.font.name = "Courier New"
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = BRAND_TEAL
para("This starts both the Vite frontend (port 5173) and FastAPI backend (port 5000) simultaneously.")

heading(2, "Option B — Start Separately")
bullet("Backend: cd backend  →  uvicorn api:app --reload --port 5000")
bullet("Frontend: npm run dev  (from project root)")

heading(2, "Environment Setup (.env file in /backend)")
add_table(
    ["Variable", "Purpose"],
    [
        ["EMAIL_HOST", "SMTP host (e.g. smtp.gmail.com)"],
        ["EMAIL_PORT", "SMTP port (usually 587)"],
        ["EMAIL_USER", "Sender email address"],
        ["EMAIL_PASS", "App password for the email account"],
        ["DATABASE_URL", "Optional — overrides default db path"],
        ["FRONTEND_ORIGIN", "Frontend URL for CORS (default: http://localhost:3000)"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  12. KNOWN LIMITATIONS & FUTURE SCOPE
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "12. Known Limitations & Future Scope")
separator()

heading(2, "Current Limitations")
bullet("OTP codes are stored in server memory — restarting the server clears all pending OTPs")
bullet("Online search (search_online.py) is a fallback with limited control over result quality")
bullet("Product/deal wizard prompts for languages other than English and Hindi fall back to English")
bullet("No user account system — identity is tied to phone/email match in the business table")

heading(2, "Future Improvements")
bullet("Persistent OTP storage in Redis or a database table")
bullet("Admin dashboard for business verification and moderation")
bullet("Push notifications for deal expiry reminders")
bullet("Google Maps integration for location-based search")
bullet("WhatsApp / Telegram bot integration")
bullet("Multi-business support for owners with more than one listing")
bullet("AI-powered business analytics (views, clicks, searches per listing)")

# Save
out_path = r"d:\Honeybee digital\chatbot frontend\frontend 4\CityHangAround_Chatbot_Documentation.docx"
doc.save(out_path)
print(f"SUCCESS! Word document saved to:\n{out_path}")
